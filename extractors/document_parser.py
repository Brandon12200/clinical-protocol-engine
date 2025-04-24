"""
Document parser for the Clinical Protocol Extraction Engine.
This module handles the conversion of various file formats (PDF, DOCX, TXT) 
to standardized plain text for processing.
"""

import os
import tempfile
import logging
import re
import json
import io
import base64
import shutil
from pathlib import Path
from collections import OrderedDict

# Document parsing libraries
import PyPDF2
import docx
import magic
import nltk
import unicodedata2 as unicodedata

# Try to import OCR-related libraries, but make them optional
try:
    import pdf2image
    from pdf2image import convert_from_path
    import pytesseract
    HAS_OCR_SUPPORT = True
except ImportError:
    HAS_OCR_SUPPORT = False
    logger.warning("OCR support disabled: pdf2image or pytesseract not available")

# Download NLTK resources if not already downloaded
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)
try:
    nltk.data.find('tokenizers/averaged_perceptron_tagger')
except LookupError:
    nltk.download('averaged_perceptron_tagger', quiet=True)

# Set up logging
logger = logging.getLogger(__name__)

class DocumentParser:
    """Parse various document formats into plain text for extraction."""
    
    # Common clinical document section headings
    CLINICAL_SECTIONS = [
        # Patient Information
        r"patient\s+information|demographics|patient\s+data",
        # Medical History
        r"medical\s+history|past\s+medical\s+history|pmh|history|past\s+history",
        # Symptoms and Presentation
        r"symptoms|chief\s+complaint|presenting\s+complaint|reason\s+for\s+visit|presentation",
        # Vital Signs
        r"vital\s+signs|vitals|observations",
        # Laboratory Results
        r"laboratory|lab\s+results|labs|test\s+results|investigations",
        # Medications
        r"medications|medication\s+list|current\s+medications|prescriptions|meds",
        # Allergies
        r"allergies|drug\s+allergies|medication\s+allergies",
        # Assessment, Diagnosis
        r"assessment|diagnosis|impression|clinical\s+impression",
        # Treatment Plan
        r"treatment|plan|management\s+plan|care\s+plan|recommendation",
        # Protocol specific
        r"inclusion\s+criteria|exclusion\s+criteria|eligibility|study\s+design|objectives",
        # Discharge Information
        r"discharge|follow[\s-]*up|discharge\s+plan|discharge\s+summary"
    ]
    
    def __init__(self, config=None):
        """
        Initialize document parser.
        
        Args:
            config (dict, optional): Configuration options for parsing.
        """
        self.config = config or {}
        self.mime_detector = magic.Magic(mime=True)
        
        # Default configuration with reasonable values
        default_config = {
            # OCR settings
            'enable_ocr': HAS_OCR_SUPPORT,
            'ocr_language': 'eng',
            'ocr_dpi': 300,
            'ocr_threshold': 10,  # Minimum text length to trigger OCR
            
            # Text processing
            'normalize_unicode': True,
            'normalize_whitespace': True,
            'remove_boilerplate': True,
            'detect_sections': True,
            
            # Table handling
            'extract_tables': True,
            'table_format': 'text'  # 'text', 'json', or 'html'
        }
        
        # Update with user config, keeping defaults for unspecified options
        for key, value in default_config.items():
            if key not in self.config:
                self.config[key] = value
                
        # Initialize NLTK tokenizer
        self.sent_tokenizer = nltk.data.load('tokenizers/punkt/english.pickle')
        
        # Compile section heading regex
        self.section_pattern = re.compile(
            r'(?:^|\n)(?P<heading>(?:{})(?:\s*:)?)(?:\n|\s*$)'.format('|'.join(self.CLINICAL_SECTIONS)), 
            re.IGNORECASE
        )
    
    def parse(self, file_path):
        """
        Parse document based on file extension.
        
        Args:
            file_path (str): Path to the document file.
            
        Returns:
            dict: Parsed document with text and metadata.
            
        Raises:
            ValueError: If file format is not supported.
            FileNotFoundError: If file does not exist.
        """
        if not os.path.exists(file_path):
            error_msg = f"File not found: {file_path}"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)
        
        file_path = os.path.abspath(file_path)
        file_ext = os.path.splitext(file_path)[1].lower()
        
        logger.info(f"Parsing document: {file_path}")
        
        try:
            # Determine parser based on file extension
            if file_ext == '.pdf':
                parsed_doc = self.parse_pdf(file_path)
            elif file_ext == '.docx':
                parsed_doc = self.parse_docx(file_path)
            elif file_ext == '.txt':
                parsed_doc = self.parse_txt(file_path)
            else:
                # Try to determine type by MIME
                mime_type = self.mime_detector.from_file(file_path)
                logger.info(f"Detected MIME type: {mime_type}")
                
                if 'pdf' in mime_type.lower():
                    parsed_doc = self.parse_pdf(file_path)
                elif 'officedocument.wordprocessingml' in mime_type.lower():
                    parsed_doc = self.parse_docx(file_path)
                elif 'text/' in mime_type.lower():
                    parsed_doc = self.parse_txt(file_path)
                else:
                    error_msg = f"Unsupported file format: {file_ext} (MIME: {mime_type})"
                    logger.error(error_msg)
                    raise ValueError(error_msg)
            
            # Add file metadata
            parsed_doc['metadata']['filename'] = os.path.basename(file_path)
            parsed_doc['metadata']['file_size'] = os.path.getsize(file_path)
            parsed_doc['metadata']['file_path'] = file_path
            parsed_doc['metadata']['file_extension'] = file_ext
            
            # Basic text stats
            parsed_doc['metadata']['char_count'] = len(parsed_doc['text'])
            parsed_doc['metadata']['word_count'] = len(parsed_doc['text'].split())
            parsed_doc['metadata']['line_count'] = parsed_doc['text'].count('\n') + 1
            
            logger.info(f"Successfully parsed document ({parsed_doc['metadata']['char_count']} chars, "
                       f"{parsed_doc['metadata']['word_count']} words)")
            
            return parsed_doc
        
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}", exc_info=True)
            raise
    
    def parse_pdf(self, file_path):
        """
        Extract text from PDF files using PyPDF2 with OCR fallback.
        
        Args:
            file_path (str): Path to PDF file.
            
        Returns:
            dict: Parsed document with text and metadata.
        """
        text = ""
        metadata = {
            'page_count': 0,
            'pages': [],
            'title': None,
            'author': None,
            'creation_date': None,
            'ocr_performed': False,
            'has_scanned_content': False
        }
        
        try:
            with open(file_path, 'rb') as file:
                pdf = PyPDF2.PdfReader(file)
                
                # Extract document metadata
                if pdf.metadata:
                    metadata['title'] = pdf.metadata.get('/Title')
                    metadata['author'] = pdf.metadata.get('/Author')
                    metadata['creation_date'] = pdf.metadata.get('/CreationDate')
                
                # Extract text from each page
                metadata['page_count'] = len(pdf.pages)
                page_texts = []
                pages_needing_ocr = []
                
                # First try with PyPDF2
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    page_text = page_text.strip()
                    
                    # Store page info
                    page_info = {
                        'page_number': i + 1,
                        'char_count': len(page_text),
                        'word_count': len(page_text.split()) if page_text else 0,
                        'ocr_performed': False
                    }
                    metadata['pages'].append(page_info)
                    
                    # Check if this page needs OCR
                    if not page_text or len(page_text) < self.config['ocr_threshold']:
                        pages_needing_ocr.append(i)
                        page_info['needs_ocr'] = True
                    else:
                        page_info['needs_ocr'] = False
                    
                    page_texts.append(page_text)
            
            # Check if we need OCR for any pages
            if pages_needing_ocr and self.config['enable_ocr'] and HAS_OCR_SUPPORT:
                logger.info(f"Performing OCR on {len(pages_needing_ocr)} pages in {file_path}")
                metadata['ocr_performed'] = True
                metadata['has_scanned_content'] = True
                
                # Create a temporary directory for the images
                with tempfile.TemporaryDirectory() as temp_dir:
                    # Convert PDF to images
                    images = convert_from_path(
                        file_path,
                        dpi=self.config['ocr_dpi'],
                        output_folder=temp_dir,
                        fmt='png',
                        output_file=f"page",
                        paths_only=True
                    )
                    
                    # Process only pages that need OCR
                    for page_idx in pages_needing_ocr:
                        if page_idx < len(images):
                            img_path = images[page_idx]
                            ocr_text = pytesseract.image_to_string(
                                img_path,
                                lang=self.config['ocr_language'],
                                config='--psm 6'  # Assume a single uniform block of text
                            )
                            
                            # Update page text and metadata
                            page_texts[page_idx] = ocr_text.strip()
                            metadata['pages'][page_idx]['char_count'] = len(ocr_text)
                            metadata['pages'][page_idx]['word_count'] = len(ocr_text.split())
                            metadata['pages'][page_idx]['ocr_performed'] = True
            
            elif pages_needing_ocr:
                if not self.config['enable_ocr']:
                    logger.warning(f"OCR disabled but needed for {len(pages_needing_ocr)} pages in {file_path}")
                elif not HAS_OCR_SUPPORT:
                    logger.warning(f"OCR not available but needed for {len(pages_needing_ocr)} pages in {file_path}")
                
                metadata['has_scanned_content'] = True
            
            # Combine all page texts
            text = "\n\n".join(page_texts)
            
            # Apply text normalization
            if self.config['normalize_unicode']:
                text = self.normalize_unicode(text)
            
            if self.config['normalize_whitespace']:
                text = self.normalize_whitespace(text)
            
            # Detect sections if enabled
            sections = []
            if self.config['detect_sections'] and text:
                sections = self.detect_sections(text)
                metadata['sections'] = sections
            
            return {
                'text': text.strip(),
                'metadata': metadata,
                'sections': sections
            }
        
        except Exception as e:
            logger.error(f"Error parsing PDF with PyPDF2/OCR: {str(e)}", exc_info=True)
            raise
    
    def normalize_unicode(self, text):
        """
        Normalize Unicode characters and handle common encoding issues.
        
        Args:
            text (str): Text to normalize.
            
        Returns:
            str: Normalized text.
        """
        if not text:
            return text
        
        # Normalize Unicode (NFKC: compatibility decomposition, followed by canonical composition)
        normalized = unicodedata.normalize('NFKC', text)
        
        # Replace common problematic characters
        replacements = {
            '\u2028': '\n',  # Line separator
            '\u2029': '\n',  # Paragraph separator
            '\u00A0': ' ',   # Non-breaking space
            '\u2013': '-',   # En dash
            '\u2014': '--',  # Em dash
            '\u2018': "'",   # Left single quotation
            '\u2019': "'",   # Right single quotation
            '\u201C': '"',   # Left double quotation
            '\u201D': '"',   # Right double quotation
            '\u2022': '•',   # Bullet
            '\u2026': '...'  # Ellipsis
        }
        
        for char, replacement in replacements.items():
            normalized = normalized.replace(char, replacement)
        
        return normalized
    
    def normalize_whitespace(self, text):
        """
        Normalize whitespace, line breaks, and remove redundant spaces.
        
        Args:
            text (str): Text to normalize.
            
        Returns:
            str: Text with normalized whitespace.
        """
        if not text:
            return text
        
        # Replace tabs with spaces
        text = text.replace('\t', ' ')
        
        # Normalize line breaks
        text = re.sub(r'\r\n?', '\n', text)
        
        # Remove multiple spaces
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove spaces at the beginning of lines
        text = re.sub(r'^ +', '', text, flags=re.MULTILINE)
        
        # Remove multiple consecutive line breaks (more than 2)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def detect_sections(self, text):
        """
        Detect common sections in clinical documents.
        
        Args:
            text (str): Document text.
            
        Returns:
            list: Identified sections with name and content.
        """
        sections = []
        
        # Find all section headings and their positions
        matches = list(self.section_pattern.finditer(text))
        
        if not matches:
            # No sections found, return whole document as one section
            return [{'name': 'Document', 'start': 0, 'end': len(text), 'content': text}]
        
        # Process each section
        for i, match in enumerate(matches):
            section_name = match.group('heading').strip()
            start_pos = match.end()
            
            # End position is the start of the next section or end of text
            if i < len(matches) - 1:
                end_pos = matches[i+1].start()
            else:
                end_pos = len(text)
            
            # Extract section content
            content = text[start_pos:end_pos].strip()
            
            # Add to sections list
            sections.append({
                'name': section_name,
                'start': start_pos,
                'end': end_pos,
                'content': content
            })
        
        return sections
    
    def parse_docx(self, file_path):
        """
        Extract text from DOCX files using python-docx.
        
        Args:
            file_path (str): Path to DOCX file.
            
        Returns:
            dict: Parsed document with text and metadata.
        """
        text = ""
        metadata = {
            'title': None,
            'author': None,
            'paragraph_count': 0,
            'has_tables': False,
            'has_images': False,
            'heading_count': 0
        }
        
        tables_data = []
        
        try:
            doc = docx.Document(file_path)
            
            # Extract document properties
            core_properties = doc.core_properties
            metadata['title'] = core_properties.title
            metadata['author'] = core_properties.author
            metadata['created'] = core_properties.created.isoformat() if core_properties.created else None
            metadata['modified'] = core_properties.modified.isoformat() if core_properties.modified else None
            
            # Track headings to help identify document structure
            headings = []
            current_heading = None
            heading_level = 0
            
            # Extract paragraphs with formatting information
            metadata['paragraph_count'] = len(doc.paragraphs)
            
            for para in doc.paragraphs:
                # Skip empty paragraphs
                if not para.text.strip():
                    continue
                
                # Check if this paragraph is a heading
                if para.style.name.startswith('Heading'):
                    try:
                        level = int(para.style.name.replace('Heading', ''))
                    except ValueError:
                        level = 1  # Default to level 1 if parsing fails
                    
                    current_heading = {
                        'text': para.text.strip(),
                        'level': level,
                        'start_position': len(text)
                    }
                    headings.append(current_heading)
                    metadata['heading_count'] += 1
                    heading_level = level
                
                # Add paragraph text
                text += para.text.strip() + "\n"
            
            # Store headings in metadata
            metadata['headings'] = headings
            
            # Check for tables
            metadata['has_tables'] = len(doc.tables) > 0
            metadata['table_count'] = len(doc.tables)
            
            # Extract tables with structure preserved
            if self.config['extract_tables']:
                for i, table in enumerate(doc.tables):
                    table_data = {
                        'id': f'table_{i+1}',
                        'rows': [],
                        'row_count': len(table.rows),
                        'col_count': len(table.rows[0].cells) if table.rows else 0
                    }
                    
                    # Process table header (first row)
                    headers = []
                    if table.rows:
                        for cell in table.rows[0].cells:
                            headers.append(cell.text.strip())
                        
                        # Process remaining rows
                        for row_idx, row in enumerate(table.rows[1:], 1):
                            row_data = {}
                            row_text = []
                            
                            for col_idx, cell in enumerate(row.cells):
                                cell_text = cell.text.strip()
                                row_text.append(cell_text)
                                
                                # If we have headers, use them as keys
                                if col_idx < len(headers) and headers[col_idx]:
                                    row_data[headers[col_idx]] = cell_text
                                else:
                                    row_data[f'col_{col_idx+1}'] = cell_text
                            
                            table_data['rows'].append(row_data)
                            
                            # Add row text to document
                            formatted_row = " | ".join(row_text)
                            if formatted_row:
                                text += formatted_row + "\n"
                    
                    tables_data.append(table_data)
                    
                    # Add table separator in document text
                    text += "\n"
            
            # Check for images (simple detection, not extraction)
            metadata['has_images'] = bool(len(doc.inline_shapes))
            metadata['image_count'] = len(doc.inline_shapes)
            
            # Apply text normalization
            if self.config['normalize_unicode']:
                text = self.normalize_unicode(text)
            
            if self.config['normalize_whitespace']:
                text = self.normalize_whitespace(text)
            
            # Detect sections if enabled
            sections = []
            if self.config['detect_sections'] and text:
                # Use both heading information and regex pattern
                sections = self.detect_sections_docx(text, headings)
                metadata['sections'] = sections
            
            return {
                'text': text.strip(),
                'metadata': metadata,
                'sections': sections,
                'tables': tables_data if tables_data else None
            }
        
        except Exception as e:
            logger.error(f"Error parsing DOCX with python-docx: {str(e)}", exc_info=True)
            raise
    
    def detect_sections_docx(self, text, headings):
        """
        Detect sections in DOCX documents using both headings and regex patterns.
        
        Args:
            text (str): Document text.
            headings (list): List of headings from DOCX document.
            
        Returns:
            list: Identified sections with name and content.
        """
        sections = []
        
        # If we have headings, use them as primary section indicators
        if headings:
            for i, heading in enumerate(headings):
                start_pos = heading['start_position']
                
                # End position is the start of the next heading or end of text
                if i < len(headings) - 1:
                    end_pos = headings[i+1]['start_position']
                else:
                    end_pos = len(text)
                
                # Extract section content
                content = text[start_pos:end_pos].strip()
                # Remove the heading text from the beginning of the content
                content = content.replace(heading['text'], '', 1).strip()
                
                # Add to sections list
                sections.append({
                    'name': heading['text'],
                    'level': heading['level'],
                    'start': start_pos,
                    'end': end_pos,
                    'content': content
                })
        
        # If no headings, fall back to regex pattern
        elif not sections:
            sections = self.detect_sections(text)
        
        return sections
    
    def parse_txt(self, file_path):
        """
        Read and parse plain text files.
        
        Args:
            file_path (str): Path to text file.
            
        Returns:
            dict: Parsed document with text and metadata.
        """
        try:
            # Detect encoding
            encoding = self.detect_encoding(file_path)
            
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()
            
            # Normalize line endings
            text = re.sub(r'\r\n?', '\n', text)
            
            metadata = {
                'encoding': encoding,
                'line_count': text.count('\n') + 1,
                'char_count': len(text),
                'word_count': len(text.split())
            }
            
            # Apply text normalization
            if self.config['normalize_unicode']:
                text = self.normalize_unicode(text)
            
            if self.config['normalize_whitespace']:
                text = self.normalize_whitespace(text)
            
            # Try to detect sections by analyzing text structure
            sections = []
            if self.config['detect_sections'] and text:
                # Look for potential section headings
                sections = self.detect_sections_txt(text)
                metadata['sections'] = sections
            
            # If clinical text, preprocess for medical terminology
            cleaned_text = self.preprocess_medical_text(text)
            
            return {
                'text': cleaned_text,
                'metadata': metadata,
                'sections': sections,
                'original_text': text if cleaned_text != text else None
            }
            
        except Exception as e:
            logger.error(f"Error parsing text file {file_path}: {str(e)}", exc_info=True)
            raise
    
    def detect_sections_txt(self, text):
        """
        Detect sections in plain text documents by looking for patterns.
        
        Args:
            text (str): Document text.
            
        Returns:
            list: Identified sections with name and content.
        """
        # First try the clinical section pattern
        sections = self.detect_sections(text)
        
        # If no clinical sections found, try to find all-caps headings 
        # (common in plain text medical documents)
        if len(sections) <= 1:
            # Look for lines that are all caps and followed by content
            caps_pattern = re.compile(r'(?:^|\n)([A-Z][A-Z\s]{3,}[A-Z0-9]:?)(?:\n|\s*$)', re.MULTILINE)
            matches = list(caps_pattern.finditer(text))
            
            if matches:
                sections = []
                for i, match in enumerate(matches):
                    section_name = match.group(1).strip()
                    start_pos = match.end()
                    
                    # End position is the start of the next section or end of text
                    if i < len(matches) - 1:
                        end_pos = matches[i+1].start()
                    else:
                        end_pos = len(text)
                    
                    # Extract section content
                    content = text[start_pos:end_pos].strip()
                    
                    # Add to sections list
                    sections.append({
                        'name': section_name,
                        'start': start_pos,
                        'end': end_pos,
                        'content': content
                    })
        
        return sections
    
    def preprocess_medical_text(self, text):
        """
        Preprocess medical text for better extraction.
        
        Args:
            text (str): Text to process.
            
        Returns:
            str: Preprocessed text.
        """
        if not text:
            return text
        
        processed_text = text
        
        # 1. Normalize common medical abbreviations
        med_abbreviations = {
            r'\bpt\b': 'patient',
            r'\bpts\b': 'patients',
            r'\bDx\b': 'diagnosis',
            r'\bRx\b': 'prescription',
            r'\bTx\b': 'treatment',
            r'\bHx\b': 'history',
            r'\bFHx\b': 'family history',
            r'\bPMH\b': 'past medical history',
            r'\bHTN\b': 'hypertension',
            r'\bDM\b': 'diabetes mellitus',
            r'\bBID\b': 'twice daily',
            r'\bTID\b': 'three times daily',
            r'\bQID\b': 'four times daily',
            r'\bPRN\b': 'as needed',
            r'\bq(\d+)h\b': r'every \1 hours',
            r'\byo\b': 'year old',
            r'\by/o\b': 'year old',
            r'\bw/\b': 'with',
            r'\bs/p\b': 'status post',
            r'\bc/o\b': 'complains of',
            r'\ba/w\b': 'associated with'
        }
        
        # Only normalize if not in the middle of a word
        for abbr, full in med_abbreviations.items():
            processed_text = re.sub(abbr, full, processed_text, flags=re.IGNORECASE)
        
        # 2. Fix spacing around measurements
        processed_text = re.sub(r'(\d+)(?:mg|mcg|g|kg|ml|mmol|mmHg|cm|mm)', r'\1 \2', processed_text)
        
        # 3. Normalize numbered lists
        processed_text = re.sub(r'(\d+)\)\s+', r'\1. ', processed_text)
        
        # 4. Normalize bullet points
        processed_text = re.sub(r'[•●◦○*]\s+', '- ', processed_text)
        
        # 5. Normalize dosage expressions
        processed_text = re.sub(r'(\d+)[-/](\d+)', r'\1/\2', processed_text)
        
        # 6. Tokenize sentences if needed (but preserve original layout)
        if self.config.get('tokenize_sentences', False):
            processed_lines = []
            for line in processed_text.split('\n'):
                if line.strip():
                    sentences = self.sent_tokenizer.tokenize(line)
                    processed_lines.append(' '.join(sentences))
                else:
                    processed_lines.append(line)
            processed_text = '\n'.join(processed_lines)
        
        return processed_text
    
    def detect_encoding(self, file_path):
        """
        Detect character encoding of text files.
        
        Args:
            file_path (str): Path to text file.
            
        Returns:
            str: Detected encoding or utf-8 as fallback.
        """
        try:
            # Use python-magic to detect mime type
            mime_type = self.mime_detector.from_file(file_path)
            
            # Check if mime type contains charset information
            if 'charset=' in mime_type:
                encoding = mime_type.split('charset=')[1].strip()
                return encoding
            
            # If no charset in mime type, try checking for BOM markers
            with open(file_path, 'rb') as file:
                raw_data = file.read(4)  # Read first 4 bytes for BOM detection
                
                if raw_data.startswith(b'\xef\xbb\xbf'):
                    return 'utf-8-sig'
                elif raw_data.startswith(b'\xff\xfe'):
                    return 'utf-16-le'
                elif raw_data.startswith(b'\xfe\xff'):
                    return 'utf-16-be'
            
            # Default to UTF-8 if no specific encoding detected
            return 'utf-8'
            
        except Exception as e:
            logger.warning(f"Error detecting encoding, falling back to utf-8: {str(e)}")
            return 'utf-8'
    
    def extract_metadata(self, file_path):
        """
        Extract metadata from document without full text extraction.
        
        Args:
            file_path (str): Path to document file.
            
        Returns:
            dict: Document metadata.
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        metadata = {
            'filename': os.path.basename(file_path),
            'file_size': os.path.getsize(file_path),
            'file_path': file_path,
            'file_extension': file_ext,
            'last_modified': os.path.getmtime(file_path),
            'mime_type': self.mime_detector.from_file(file_path)
        }
        
        try:
            if file_ext == '.pdf':
                # Extract just PDF metadata
                with open(file_path, 'rb') as file:
                    pdf = PyPDF2.PdfReader(file)
                    if pdf.metadata:
                        metadata.update({
                            'title': pdf.metadata.get('/Title'),
                            'author': pdf.metadata.get('/Author'),
                            'creation_date': pdf.metadata.get('/CreationDate'),
                            'page_count': len(pdf.pages)
                        })
            
            elif file_ext == '.docx':
                # Extract just DOCX metadata
                doc = docx.Document(file_path)
                core_properties = doc.core_properties
                metadata.update({
                    'title': core_properties.title,
                    'author': core_properties.author,
                    'created': core_properties.created.isoformat() if core_properties.created else None,
                    'modified': core_properties.modified.isoformat() if core_properties.modified else None,
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables),
                    'image_count': len(doc.inline_shapes)
                })
            
        except Exception as e:
            logger.warning(f"Error extracting detailed metadata: {str(e)}")
        
        return metadata
    
    def get_supported_formats(self):
        """Return list of supported document formats."""
        return ['.pdf', '.docx', '.txt']


# For easy testing
if __name__ == "__main__":
    import argparse
    import json
    
    # Configure logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    parser = argparse.ArgumentParser(description='Parse document to text')
    parser.add_argument('file_path', help='Path to document file')
    parser.add_argument('--metadata-only', action='store_true', help='Extract only metadata')
    parser.add_argument('--output', help='Output file for extracted text')
    args = parser.parse_args()
    
    try:
        parser = DocumentParser()
        
        if args.metadata_only:
            result = parser.extract_metadata(args.file_path)
            print(json.dumps(result, indent=2))
        else:
            result = parser.parse(args.file_path)
            
            # Print metadata
            print("\nDocument Metadata:")
            print(json.dumps(result['metadata'], indent=2))
            
            # Handle output
            if args.output:
                with open(args.output, 'w', encoding='utf-8') as f:
                    f.write(result['text'])
                print(f"\nText saved to {args.output}")
            else:
                # Print preview of text
                text_preview = result['text'][:500]
                if len(result['text']) > 500:
                    text_preview += "..."
                print("\nText Preview:")
                print(text_preview)
                print(f"\nFull text length: {len(result['text'])} characters")
    
    except Exception as e:
        logger.error(f"Error: {str(e)}", exc_info=True)
        print(f"\nError: {str(e)}")